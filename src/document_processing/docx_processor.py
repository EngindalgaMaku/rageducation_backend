"""
DOCX document processing module.

This module provides functions to extract text from DOCX files using python-docx.
"""

import docx
from ..utils.helpers import setup_logging

logger = setup_logging()

def process_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file.

    Args:
        file_path: The path to the DOCX file.

    Returns:
        The extracted text as a single string, or an empty string if an error occurs.
    """
    logger.info(f"Processing DOCX file: {file_path}")
    text = ""
    try:
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text += para.text + "\n"
        logger.info(f"Successfully extracted text from {file_path}")
        return text.strip()
    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return ""
    except Exception as e:
        logger.error(f"An error occurred while processing {file_path}: {e}")
        return ""

if __name__ == '__main__':
    # This is for testing purposes.
    # Create a dummy DOCX file to test the processor.
    import os
    if not os.path.exists('data/raw/sample_documents'):
        os.makedirs('data/raw/sample_documents')

    test_docx_path = "data/raw/sample_documents/sample.docx"
    
    # Create a dummy docx file for testing
    doc = docx.Document()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It contains multiple paragraphs for testing the RAG system.")
    doc.save(test_docx_path)

    if os.path.exists(test_docx_path):
        extracted_text = process_docx(test_docx_path)
        if extracted_text:
            print("--- Extracted Text from DOCX ---")
            print(extracted_text)
            print("------------------------------")
        else:
            print("Could not extract text from the DOCX.")
    else:
        print(f"Test file not found: {test_docx_path}. Please create it to run the test.")
